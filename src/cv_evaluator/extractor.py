from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Optional


@dataclass
class ExtractedText:
    text: str
    pages: int | None = None
    file_type: str | None = None
    notes: list[str] | None = None


def _read_txt(path: Path) -> ExtractedText:
    # Try utf-8 then fallback to binary detect
    try:
        t = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        import chardet

        raw = path.read_bytes()
        enc = chardet.detect(raw).get("encoding") or "utf-8"
        t = raw.decode(enc, errors="ignore")
    return ExtractedText(text=t, pages=None, file_type="txt", notes=[])


def _read_docx(path: Path) -> ExtractedText:
    from docx import Document

    doc = Document(path)
    parts: list[str] = []
    notes: list[str] = []
    table_count = len(doc.tables)
    if table_count:
        notes.append(f"contains {table_count} table(s)")
    # Image detection (heuristic): count inline shapes
    try:
        img_count = len(getattr(doc, "inline_shapes", []))
        if img_count:
            notes.append(f"contains {img_count} image(s)")
    except Exception:
        pass
    for p in doc.paragraphs:
        parts.append(p.text)
    # Very rough page estimate based on paragraphs
    pages = max(1, (len(parts) // 40)) if parts else 1
    return ExtractedText(text="\n".join(parts), pages=pages, file_type="docx", notes=notes)


def _read_pdf(path: Path) -> ExtractedText:
    # First try PyPDF2 for page count; pdfminer for text extraction
    from PyPDF2 import PdfReader
    from pdfminer.high_level import extract_text as pdf_extract_text

    notes: list[str] = []
    pages: Optional[int] = None
    reader = None
    try:
        reader = PdfReader(str(path))
        pages = len(reader.pages)
    except Exception:
        notes.append("could not read page count with PyPDF2")
    try:
        txt = pdf_extract_text(str(path)) or ""
    except Exception:
        txt = ""
        notes.append("failed to extract text with pdfminer")
    # Heuristic image presence: scan XObjects for /Image
    if reader is not None:
        try:
            img_total = 0
            for page in reader.pages:
                resources = page.get("/Resources") or {}
                xobj = resources.get("/XObject") if hasattr(resources, "get") else None
                if xobj:
                    # xobj may be an IndirectObject mapping
                    xobj = xobj.get_object()
                    for obj in xobj.values():
                        try:
                            o = obj.get_object()
                            if o.get("/Subtype") == "/Image":
                                img_total += 1
                        except Exception:
                            continue
            if img_total:
                notes.append(f"contains {img_total} image(s)")
        except Exception:
            pass
    return ExtractedText(text=txt, pages=pages, file_type="pdf", notes=notes)


def extract_text(path: Path) -> ExtractedText:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return _read_txt(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    # Fallback: read as text
    return _read_txt(path)
